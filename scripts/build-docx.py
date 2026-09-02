#!/usr/bin/env python3
r"""Build a Word (.docx) version of the CV from its LaTeX content file.

    scripts/build-docx.py CV-EN/main.tex dist/Luca_Tomei_CV_EN.docx

The content files only use the small macro set defined in
common/preamble.tex (\cvheader, \section, \cvskill, \cventry, \cvsub,
\cvstack, \cvline, \cvproject, \cvcert, \cvlanguage, \cvprivacy and the
cvitems list), so this script walks those macros and emits real Word
structures: heading styles, bullet lists, right-aligned dates. Some job
portals only accept Word files, and a native .docx parses better than a
PDF converted by hand.

Requires python-docx.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
MUTED = RGBColor(0x5B, 0x64, 0x70)
FONT = "Calibri"

# --------------------------------------------------------------- LaTeX text

ACCENTS = {"`": "\u0300", "'": "\u0301", "^": "\u0302", '"': "\u0308", "~": "\u0303"}


def tex_to_text(s: str) -> str:
    """Turn a fragment of the CV's LaTeX into plain text."""
    import unicodedata

    s = re.sub(r"(?<!\\)%.*", "", s)                       # comments
    s = re.sub(r"\\href\{[^}]*\}\{((?:[^{}]|\{[^{}]*\})*)\}", r"\1", s)
    for cmd in ("textbf", "textit", "emph", r"textcolor\{[a-z]+\}"):
        s = re.sub(r"\\" + cmd + r"\{((?:[^{}]|\{[^{}]*\})*)\}", r"\1", s)
    s = re.sub(r"\\cvlanguage\{([^}]*)\}\{([^}]*)\}", r"\1 – \2", s)
    s = re.sub(r"\\([`'^\"~])([a-zA-Z])", lambda m: m.group(2) + ACCENTS[m.group(1)], s)
    s = s.replace("\\cvsep\\", " · ").replace("\\cvsep", " · ")
    s = s.replace("\\textendash\\", "–").replace("\\textendash", "–")
    s = s.replace("\\textemdash\\", "—").replace("\\textemdash", "—")
    s = s.replace("\\textperiodcentered", "·").replace("\\textbullet", "•")
    s = s.replace("\\par", "\x00").replace("\\\\", "\x00")   # explicit breaks only
    s = re.sub(r"\s*\n\s*", " ", s)                       # source line wraps are spaces
    s = re.sub(r"\\([#&%_$])", r"\1", s)
    s = s.replace("\\ ", " ").replace("~", " ")
    s = re.sub(r"(?<!-)---(?!-)", "—", s)
    s = re.sub(r"(?<!-)--(?!-)", "–", s)
    s = re.sub(r"\\[a-zA-Z]+\*?", "", s)                  # leftover commands
    s = s.replace("{", "").replace("}", "")
    s = unicodedata.normalize("NFC", s)
    s = re.sub(r"[ \t]+", " ", s)
    return "\n".join(line.strip() for line in s.split("\x00")).strip()


def read_args(src: str, pos: int, n: int):
    """Read n brace groups starting at pos; return (args, end_pos)."""
    args = []
    for _ in range(n):
        while src[pos].isspace():
            pos += 1
        assert src[pos] == "{", src[pos:pos + 40]
        depth, start = 0, pos
        while True:
            c = src[pos]
            if c == "{":
                depth += 1
            elif c == "}":
                depth -= 1
                if depth == 0:
                    break
            pos += 1
        args.append(src[start + 1:pos])
        pos += 1
    return args, pos


MACROS = {
    "cvheader": 3, "section": 1, "cvskill": 2, "cventry": 4, "cvsub": 2,
    "cvstack": 2, "cvline": 3, "cvproject": 3, "cvcert": 3, "cvprivacy": 1,
}


def parse(src: str):
    """Yield (kind, payload) events in document order."""
    body = src[src.index(r"\begin{document}") + len(r"\begin{document}"):src.index(r"\end{document}")]
    pos = 0
    text_buf = []

    def flush():
        txt = tex_to_text("\n".join(text_buf))
        text_buf.clear()
        if txt:
            yield ("text", txt)

    pattern = re.compile(r"\\(" + "|".join(MACROS) + r"|cvlanguage)(?![a-zA-Z])|\\(begin\{cvitems\})")
    while True:
        m = pattern.search(body, pos)
        if not m:
            text_buf.append(body[pos:])
            yield from flush()
            break
        text_buf.append(body[pos:m.start()])
        name = m.group(1) or m.group(2)
        if name == "begin{cvitems}":
            yield from flush()
            end = body.index(r"\end{cvitems}", m.end())
            items = re.split(r"\\item\s", body[m.end():end])
            yield ("items", [tex_to_text(i) for i in items if i.strip()])
            pos = end + len(r"\end{cvitems}")
        elif name == "cvlanguage":
            # languages sit on one line, possibly several per line
            line_end = body.index("\n", m.start())
            yield from flush()
            yield ("text", tex_to_text(body[m.start():line_end]))
            pos = line_end
        else:
            yield from flush()
            args, pos = read_args(body, m.end(), MACROS[name])
            yield (name, [tex_to_text(a) for a in args])


# ------------------------------------------------------------------- Word

def set_cell_font(run, size=10.5, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def para(doc, text="", size=10.5, bold=False, italic=False, color=None,
         style=None, before=0, after=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if text:
        set_cell_font(p.add_run(text), size, bold, italic, color)
    return p


def title_date(doc, left, right, bold=True, size=10.5, before=4):
    """Bold title on the left, muted date flush right (via a right tab)."""
    p = para(doc, before=before)
    section = doc.sections[0]
    width = section.page_width - section.left_margin - section.right_margin
    p.paragraph_format.tab_stops.add_tab_stop(width, WD_TAB_ALIGNMENT.RIGHT)
    set_cell_font(p.add_run(left), size, bold=bold)
    if right:
        set_cell_font(p.add_run("\t" + right), size, color=MUTED)
    return p


def heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_cell_font(r, 13, bold=True, color=ACCENT)
    # bottom border under the heading
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "C9CED6")):
        bottom.set(qn(k), v)
    pbdr.append(bottom)
    pPr.append(pbdr)


def build(tex_path: Path, out_path: Path, author="Luca Tomei"):
    src = tex_path.read_text(encoding="utf-8")
    lang = re.search(r"\\newcommand\{\\cvlang\}\{([^}]*)\}", src).group(1)
    title = tex_to_text(re.search(r"\\newcommand\{\\cvpdftitle\}\{([^}]*)\}", src).group(1))
    subject = tex_to_text(re.search(r"\\newcommand\{\\cvpdfsubject\}\{([^}]*)\}", src).group(1))
    keywords = tex_to_text(re.search(r"\\newcommand\{\\cvpdfkeywords\}\{((?:[^{}]|\{[^{}]*\})*)\}", src).group(1))

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    for side in ("left_margin", "right_margin"):
        setattr(sec, side, Cm(1.6))
    sec.top_margin = sec.bottom_margin = Cm(1.4)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    doc.core_properties.author = author
    doc.core_properties.title = title
    doc.core_properties.subject = subject
    doc.core_properties.keywords = keywords[:255]   # Word's limit
    doc.core_properties.language = lang

    for kind, payload in parse(src):
        if kind == "cvheader":
            name, headline, contacts = payload
            para(doc, name, size=24, bold=True, color=ACCENT)
            para(doc, headline, size=12, color=MUTED)
            for line in contacts.split("\n"):
                para(doc, line, size=10, color=MUTED)
            para(doc, after=2)
        elif kind == "section":
            heading(doc, payload[0])
        elif kind == "text":
            for line in payload.split("\n"):
                para(doc, line, after=2)
        elif kind == "cvskill":
            p = para(doc, after=1)
            set_cell_font(p.add_run(payload[0] + ": "), bold=True, color=ACCENT)
            set_cell_font(p.add_run(payload[1]))
        elif kind == "cventry":
            t, d, org, loc = payload
            title_date(doc, t, d)
            para(doc, org + (" – " + loc if loc else ""), italic=True, color=MUTED)
        elif kind == "cvsub":
            title_date(doc, payload[0], payload[1])
        elif kind == "cvline":
            title_date(doc, payload[0] + " – " + payload[2], payload[1])
        elif kind == "cvstack":
            p = para(doc, before=1, after=2)
            set_cell_font(p.add_run(payload[0] + ": "), size=9.5, italic=True)
            set_cell_font(p.add_run(payload[1]), size=9.5)
        elif kind == "cvproject":
            n, right, url = payload
            title_date(doc, n + (" – " + url if url else ""), right, before=3)
        elif kind == "cvcert":
            n, issuer, year = payload
            p = para(doc, after=1)
            set_cell_font(p.add_run(n), size=10, bold=True)
            set_cell_font(p.add_run(" – " + issuer + (", " + year if year else "")), size=10)
        elif kind == "items":
            for it in payload:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                set_cell_font(p.add_run(it))
        elif kind == "cvprivacy":
            para(doc, payload[0], size=8.5, italic=True, color=MUTED, before=8)
            para(doc, author, before=6, align=WD_ALIGN_PARAGRAPH.RIGHT)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        sys.exit(__doc__)
    build(Path(sys.argv[1]), Path(sys.argv[2]))
